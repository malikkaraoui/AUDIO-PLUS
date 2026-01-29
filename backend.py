"""
AUDIO PLUS — API Flask pour découper et transcrire des MP3.
"""

import json
import os
import tempfile
import threading
import requests as http_requests
from flask import Flask, request, jsonify, send_file, Response
from flask_cors import CORS
from pydub import AudioSegment
import whisper
from docx import Document

app = Flask(__name__)
CORS(app)

# Cancel flag for stopping transcription
cancel_flag = threading.Event()

SEGMENT_DURATION_MS = 10 * 60 * 1000  # 10 minutes
CHUNK_DURATION_MS = 5 * 60 * 1000  # 5 min chunks for progress tracking
OUTPUT_DIR = os.path.expanduser("~/Downloads")

whisper_model = None


def get_model():
    global whisper_model
    if whisper_model is None:
        whisper_model = whisper.load_model("medium")
    return whisper_model


def format_duration(ms):
    """Format milliseconds as MM:SS."""
    s = int(ms / 1000)
    return f"{s // 60}:{s % 60:02d}"


def transcribe_with_progress(file_path, use_cleanup=True, ollama_model="llama3.1:8b"):
    """Split audio into chunks, transcribe each, yield progress."""
    audio = AudioSegment.from_mp3(file_path)
    total = len(audio)
    chunks = list(range(0, total, CHUNK_DURATION_MS))
    n = len(chunks)
    model = get_model()
    texts = []
    total_steps = n * 2 if use_cleanup else n

    cancel_flag.clear()
    yield 0, "", "init", f"Audio chargé — durée {format_duration(total)}, {n} segment(s) de 5 min"

    for i, start in enumerate(chunks):
        if cancel_flag.is_set():
            yield 0, "", "cancelled", "Annulé par l'utilisateur"
            return
        end = min(start + CHUNK_DURATION_MS, total)
        log = f"[{i+1}/{n}] Whisper — segment {format_duration(start)}→{format_duration(end)}"
        seg = audio[start : start + CHUNK_DURATION_MS]
        tmp = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False)
        seg.export(tmp.name, format="mp3")
        tmp.close()

        try:
            result = model.transcribe(tmp.name, language="fr")
            raw_text = result["text"]
        finally:
            os.unlink(tmp.name)

        step = (i + 1)
        progress = int(step / total_steps * 100)
        yield progress, " ".join(texts + [raw_text]), "transcription", f"{log} ✓"

        if use_cleanup:
            if cancel_flag.is_set():
                yield 0, "", "cancelled", "Annulé par l'utilisateur"
                return
            log_c = f"[{i+1}/{n}] Correction Ollama ({ollama_model})"
            yield progress, " ".join(texts + [raw_text]), "correction", log_c
            cleaned = cleanup_with_ollama(raw_text, model=ollama_model)
            texts.append(cleaned)
            step = n + (i + 1)
            progress = int(step / total_steps * 100)
            yield progress, " ".join(texts), "correction", f"{log_c} ✓"
        else:
            texts.append(raw_text)

    return " ".join(texts)


OLLAMA_BASE = "http://localhost:11434"
OLLAMA_URL = OLLAMA_BASE + "/api/generate"

CLEANUP_PROMPT = """Tu es un correcteur de transcription audio. Réponds UNIQUEMENT en français.
Voici une transcription brute. Corrige UNIQUEMENT :
- La ponctuation (points, virgules, points d'interrogation)
- Les majuscules en début de phrase
- Les fautes d'orthographe évidentes

RÈGLES STRICTES :
- NE CHANGE PAS les mots, les phrases, l'ordre, le sens ou le style.
- NE RÉSUME PAS. NE REFORMULE PAS. NE RAJOUTE RIEN.
- N'INVENTE AUCUN MOT ni AUCUNE PHRASE qui n'existe pas dans l'original.
- Garde le texte le plus fidèle possible à l'original.
- Réponds UNIQUEMENT avec le texte corrigé, rien d'autre.

Transcription brute :
"""


def cleanup_with_ollama(text, model="llama3.1:8b"):
    """Clean up transcription using local Ollama LLM."""
    try:
        resp = http_requests.post(OLLAMA_URL, json={
            "model": model,
            "prompt": CLEANUP_PROMPT + text,
            "stream": False,
        }, timeout=300)
        if resp.ok:
            return resp.json().get("response", text)
    except Exception:
        pass
    return text


@app.route("/models", methods=["GET"])
def models():
    try:
        resp = http_requests.get(OLLAMA_BASE + "/api/tags", timeout=5)
        if resp.ok:
            names = [m["name"] for m in resp.json().get("models", [])]
            return jsonify({"models": names})
    except Exception:
        pass
    return jsonify({"models": []})


@app.route("/stop", methods=["POST"])
def stop():
    cancel_flag.set()
    return jsonify({"stopped": True})


@app.route("/split", methods=["POST"])
def split():
    file = request.files.get("file")
    if not file or not file.filename.lower().endswith(".mp3"):
        return jsonify({"error": "Fichier MP3 requis"}), 400

    base_name = os.path.splitext(file.filename)[0]
    audio = AudioSegment.from_mp3(file)
    total = len(audio)
    parts = []

    for i, start in enumerate(range(0, total, SEGMENT_DURATION_MS)):
        segment = audio[start : start + SEGMENT_DURATION_MS]
        out_name = f"{base_name} - partie {i + 1}.mp3"
        out_path = os.path.join(OUTPUT_DIR, out_name)
        segment.export(out_path, format="mp3")
        parts.append(out_name)

    return jsonify({"parts": parts, "output_dir": OUTPUT_DIR})


@app.route("/transcribe", methods=["POST"])
def transcribe():
    file = request.files.get("file")
    if not file or not file.filename.lower().endswith(".mp3"):
        return jsonify({"error": "Fichier MP3 requis"}), 400

    ollama_model = request.form.get("model", "llama3.1:8b")
    tmp = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False)
    file.save(tmp.name)
    tmp.close()

    def generate():
        try:
            for progress, text, phase, log in transcribe_with_progress(tmp.name, ollama_model=ollama_model):
                yield f"data: {json.dumps({'progress': progress, 'text': text, 'phase': phase, 'log': log})}\n\n"
        finally:
            os.unlink(tmp.name)

    return Response(generate(), mimetype="text/event-stream")


@app.route("/transcribe/docx", methods=["POST"])
def transcribe_docx():
    file = request.files.get("file")
    if not file or not file.filename.lower().endswith(".mp3"):
        return jsonify({"error": "Fichier MP3 requis"}), 400

    base_name = os.path.splitext(file.filename)[0]
    ollama_model = request.form.get("model", "llama3.1:8b")

    tmp = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False)
    file.save(tmp.name)
    tmp.close()

    def generate():
        try:
            text = ""
            for progress, text, phase, log in transcribe_with_progress(tmp.name, ollama_model=ollama_model):
                yield f"data: {json.dumps({'progress': progress, 'phase': phase, 'log': log})}\n\n"

            # Create DOCX
            doc = Document()
            doc.add_heading(base_name, level=1)
            doc.add_paragraph(text)
            docx_path = os.path.join(OUTPUT_DIR, f"{base_name} - transcription.docx")
            doc.save(docx_path)

            yield f"data: {json.dumps({'progress': 100, 'done': True, 'path': docx_path})}\n\n"
        finally:
            os.unlink(tmp.name)

    return Response(generate(), mimetype="text/event-stream")


if __name__ == "__main__":
    app.run(port=5001, debug=True)
